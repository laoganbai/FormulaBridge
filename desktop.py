import sys
from pathlib import Path

import webview
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
PUBLIC_DIR = BASE_DIR / "public"


def safe_name(value):
    cleaned = "".join(character for character in str(value or "公式解析结果") if character not in '<>:"/\\|?*')
    cleaned = " ".join(cleaned.split()).strip()[:80]
    return cleaned or "公式解析结果"


def math_element(name):
    return OxmlElement(f"m:{name}")


def math_text(value, bold=False):
    run = math_element("r")
    if bold:
        properties = math_element("rPr")
        style = math_element("sty")
        style.set(qn("m:val"), "b")
        properties.append(style)
        run.append(properties)
    text_node = math_element("t")
    text_node.text = str(value or "")
    run.append(text_node)
    return run


def node_list(value):
    if isinstance(value, list):
        return value
    return [value] if value else []


def plain_ast(nodes):
    parts = []
    for node in node_list(nodes):
        if not isinstance(node, dict):
            continue
        kind = node.get("type")
        if kind == "text":
            parts.append(str(node.get("value", "")))
        elif kind in {"group", "styled", "accent"}:
            parts.append(plain_ast(node.get("children")))
        elif kind == "frac":
            parts.append(f"{plain_ast(node.get('numerator'))}/{plain_ast(node.get('denominator'))}")
        elif kind == "sqrt":
            parts.append(f"√{plain_ast(node.get('radicand'))}")
        elif kind == "sup":
            parts.append(f"{plain_ast(node.get('base'))}{plain_ast(node.get('sup'))}")
        elif kind == "sub":
            parts.append(f"{plain_ast(node.get('base'))}{plain_ast(node.get('sub'))}")
        elif kind == "subsup":
            parts.append(f"{plain_ast(node.get('base'))}{plain_ast(node.get('sub'))}{plain_ast(node.get('sup'))}")
        elif kind == "binom":
            parts.append(f"({plain_ast(node.get('numerator'))}/{plain_ast(node.get('denominator'))})")
        elif kind == "matrix":
            rows = []
            for row in node.get("rows", []):
                rows.append(", ".join(plain_ast(cell) for cell in row))
            parts.append("[" + "; ".join(rows) + "]")
    return "".join(parts)


def append_math(parent, nodes, bold=False):
    for node in node_list(nodes):
        if not isinstance(node, dict):
            continue
        kind = node.get("type")
        if kind == "text":
            parent.append(math_text(node.get("value", ""), bold=bold))
        elif kind in {"group", "accent"}:
            append_math(parent, node.get("children"), bold=bold)
        elif kind == "styled":
            append_math(parent, node.get("children"), bold=bold or node.get("style") == "bold")
        elif kind == "frac":
            fraction = math_element("f")
            numerator = math_element("num")
            denominator = math_element("den")
            append_math(numerator, node.get("numerator"), bold=bold)
            append_math(denominator, node.get("denominator"), bold=bold)
            fraction.extend([numerator, denominator])
            parent.append(fraction)
        elif kind == "sqrt":
            radical = math_element("rad")
            if node.get("degree"):
                degree = math_element("deg")
                append_math(degree, node.get("degree"), bold=bold)
                radical.append(degree)
            radicand = math_element("e")
            append_math(radicand, node.get("radicand"), bold=bold)
            radical.append(radicand)
            parent.append(radical)
        elif kind in {"sup", "sub", "subsup"}:
            container = math_element({"sup": "sSup", "sub": "sSub", "subsup": "sSubSup"}[kind])
            base = math_element("e")
            append_math(base, node.get("base"), bold=bold)
            container.append(base)
            if kind in {"sub", "subsup"}:
                sub = math_element("sub")
                append_math(sub, node.get("sub"), bold=bold)
                container.append(sub)
            if kind in {"sup", "subsup"}:
                sup = math_element("sup")
                append_math(sup, node.get("sup"), bold=bold)
                container.append(sup)
            parent.append(container)
        elif kind == "binom":
            parent.append(math_text("(", bold=bold))
            append_math(parent, {"type": "frac", "numerator": node.get("numerator"), "denominator": node.get("denominator")}, bold=bold)
            parent.append(math_text(")", bold=bold))
        elif kind == "matrix":
            parent.append(math_text(plain_ast(node), bold=bold))


def add_math_paragraph(document, item):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_root = math_element("oMath")
    append_math(math_root, item.get("ast"))
    if len(math_root) == 0:
        paragraph.add_run(item.get("source", ""))
    else:
        paragraph._p.append(math_root)
    return paragraph


def add_source_paragraph(document, source):
    paragraph = document.add_paragraph()
    label = paragraph.add_run("源代码  ")
    label.bold = True
    label.font.size = Pt(9)
    label.font.color.rgb = RGBColor(97, 115, 107)
    source_run = paragraph.add_run(str(source or ""))
    source_run.font.name = "Cascadia Mono"
    source_run.font.size = Pt(9)
    source_run.font.color.rgb = RGBColor(97, 115, 107)
    return paragraph


def create_docx(payload, target):
    title = str(payload.get("title") or "AI 公式解析结果").strip()[:120] or "AI 公式解析结果"
    items = payload.get("items") if isinstance(payload.get("items"), list) else []
    document = Document()
    document.core_properties.title = title
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.right_margin = Inches(0.8)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)

    heading = document.add_paragraph()
    heading.style = document.styles["Title"]
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.name = "Aptos Display"
    heading_run.font.size = Pt(18)
    heading_run.font.color.rgb = RGBColor(23, 52, 42)

    meta = document.add_paragraph(f"{sum(item.get('type') == 'formula' for item in items)} 个公式")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(97, 115, 107)

    for item in items:
        kind = item.get("type")
        if kind == "formula":
            add_math_paragraph(document, item)
            if payload.get("includeSource", True):
                add_source_paragraph(document, item.get("source", ""))
        elif kind in {"text", "code"}:
            for line in str(item.get("text", "")).splitlines() or [""]:
                paragraph = document.add_paragraph(line or " ")
                if kind == "code" and paragraph.runs:
                    paragraph.runs[0].font.name = "Cascadia Mono"
                    paragraph.runs[0].font.size = Pt(9)

    if not items:
        document.add_paragraph("未检测到可导出的内容。")
    document.save(str(target))


class DesktopApi:
    def export_docx(self, payload):
        payload = payload if isinstance(payload, dict) else {}
        filename = safe_name(payload.get("fileName") or payload.get("title")) + ".docx"
        documents = Path.home() / "Documents"
        documents.mkdir(parents=True, exist_ok=True)
        selected = webview.windows[0].create_file_dialog(
            webview.SAVE_DIALOG,
            directory=str(documents),
            save_filename=filename,
            file_types=("Word document (*.docx)",),
        )
        if isinstance(selected, (tuple, list)):
            selected = selected[0] if selected else None
        if not selected:
            return {"canceled": True}
        target = Path(selected)
        if target.suffix.lower() != ".docx":
            target = target.with_suffix(".docx")
        create_docx(payload, target)
        return {"canceled": False, "filePath": target.name}


if __name__ == "__main__":
    window = webview.create_window(
        "公式解析器",
        str(PUBLIC_DIR / "index.html"),
        js_api=DesktopApi(),
        width=1360,
        height=900,
        min_size=(960, 680),
        background_color="#f2f5f1",
    )
    webview.start(debug=False)
